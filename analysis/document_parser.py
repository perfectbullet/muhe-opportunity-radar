"""
文档解析器模块
支持 PDF、Word、Markdown 文档的文本提取
"""

from pathlib import Path
from typing import Dict, Optional, Union
import logging

# 配置日志
logger = logging.getLogger(__name__)


class DocumentParser:
    """文档解析器 - 支持多种格式"""
    
    SUPPORTED_FORMATS = {
        'pdf': ['.pdf'],
        'word': ['.doc', '.docx'],
        'markdown': ['.md', '.markdown']
    }
    
    def __init__(self):
        """初始化文档解析器"""
        self._check_dependencies()
    
    def _check_dependencies(self):
        """检查并记录可用的解析库"""
        self.available_parsers = {}
        
        # 检查 PDF 解析库
        try:
            import pdfplumber
            self.available_parsers['pdf'] = 'pdfplumber'
            logger.info("✓ pdfplumber 可用")
        except ImportError:
            try:
                import PyPDF2
                self.available_parsers['pdf'] = 'pypdf2'
                logger.info("✓ PyPDF2 可用")
            except ImportError:
                logger.warning("⚠️  PDF 解析库未安装 (pdfplumber 或 PyPDF2)")
        
        # 检查 Word 解析库
        try:
            import docx
            self.available_parsers['word'] = 'python-docx'
            logger.info("✓ python-docx 可用")
        except ImportError:
            logger.warning("⚠️  Word 解析库未安装 (python-docx)")
        
        # Markdown 无需额外依赖
        self.available_parsers['markdown'] = 'builtin'
    
    def parse(self, file_path: Union[str, Path]) -> Dict[str, any]:
        """
        解析文档并提取文本
        
        Args:
            file_path: 文档文件路径
            
        Returns:
            包含以下字段的字典:
            - content: 提取的文本内容
            - format: 文档格式 (pdf/word/markdown)
            - pages: 页数（如适用）
            - metadata: 元数据信息
            - success: 是否成功解析
            - error: 错误信息（如失败）
        """
        file_path = Path(file_path)
        
        # 检查文件是否存在
        if not file_path.exists():
            return {
                "success": False,
                "error": f"文件不存在: {file_path}",
                "content": None,
                "format": None
            }
        
        # 识别文件格式
        file_format = self._identify_format(file_path)
        if not file_format:
            return {
                "success": False,
                "error": f"不支持的文件格式: {file_path.suffix}",
                "content": None,
                "format": None
            }
        
        # 根据格式调用相应的解析方法
        try:
            if file_format == 'pdf':
                result = self._parse_pdf(file_path)
            elif file_format == 'word':
                result = self._parse_word(file_path)
            elif file_format == 'markdown':
                result = self._parse_markdown(file_path)
            else:
                return {
                    "success": False,
                    "error": f"未实现的解析器: {file_format}",
                    "content": None,
                    "format": file_format
                }
            
            result['format'] = file_format
            result['success'] = True
            return result
            
        except Exception as e:
            logger.error(f"解析文档失败 {file_path}: {str(e)}")
            return {
                "success": False,
                "error": str(e),
                "content": None,
                "format": file_format
            }
    
    def _identify_format(self, file_path: Path) -> Optional[str]:
        """识别文件格式"""
        suffix = file_path.suffix.lower()
        
        for format_name, extensions in self.SUPPORTED_FORMATS.items():
            if suffix in extensions:
                return format_name
        
        return None
    
    def _parse_pdf(self, file_path: Path) -> Dict:
        """解析 PDF 文档"""
        parser = self.available_parsers.get('pdf')
        
        if not parser:
            raise ImportError("PDF 解析库未安装，请运行: pip install pdfplumber")
        
        if parser == 'pdfplumber':
            return self._parse_pdf_pdfplumber(file_path)
        elif parser == 'pypdf2':
            return self._parse_pdf_pypdf2(file_path)
    
    def _parse_pdf_pdfplumber(self, file_path: Path) -> Dict:
        """使用 pdfplumber 解析 PDF"""
        import pdfplumber
        
        text_content = []
        tables = []
        metadata = {}
        
        with pdfplumber.open(file_path) as pdf:
            metadata = {
                'pages': len(pdf.pages),
                'metadata': pdf.metadata
            }
            
            for i, page in enumerate(pdf.pages):
                # 提取文本
                page_text = page.extract_text()
                if page_text:
                    text_content.append(f"--- 第 {i+1} 页 ---\n{page_text}")
                
                # 提取表格（可选）
                page_tables = page.extract_tables()
                if page_tables:
                    tables.extend(page_tables)
        
        return {
            "content": "\n\n".join(text_content),
            "pages": metadata['pages'],
            "metadata": metadata,
            "tables": tables if tables else None
        }
    
    def _parse_pdf_pypdf2(self, file_path: Path) -> Dict:
        """使用 PyPDF2 解析 PDF（备用方案）"""
        import PyPDF2
        
        text_content = []
        
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            num_pages = len(pdf_reader.pages)
            
            for i, page in enumerate(pdf_reader.pages):
                text = page.extract_text()
                if text:
                    text_content.append(f"--- 第 {i+1} 页 ---\n{text}")
            
            metadata = pdf_reader.metadata if pdf_reader.metadata else {}
        
        return {
            "content": "\n\n".join(text_content),
            "pages": num_pages,
            "metadata": dict(metadata)
        }
    
    def _parse_word(self, file_path: Path) -> Dict:
        """解析 Word 文档"""
        if 'word' not in self.available_parsers:
            raise ImportError("Word 解析库未安装，请运行: pip install python-docx")
        
        from docx import Document
        
        doc = Document(file_path)
        
        # 提取段落文本
        paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
        
        # 提取表格（可选）
        tables = []
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text for cell in row.cells]
                table_data.append(row_data)
            tables.append(table_data)
        
        # 元数据
        metadata = {
            'paragraphs': len(paragraphs),
            'tables': len(tables),
            'core_properties': {
                'author': doc.core_properties.author,
                'created': str(doc.core_properties.created),
                'modified': str(doc.core_properties.modified),
                'title': doc.core_properties.title
            }
        }
        
        return {
            "content": "\n\n".join(paragraphs),
            "metadata": metadata,
            "tables": tables if tables else None
        }
    
    def _parse_markdown(self, file_path: Path) -> Dict:
        """解析 Markdown 文档"""
        with open(file_path, 'r', encoding='utf-8') as file:
            content = file.read()
        
        # 统计基本信息
        lines = content.split('\n')
        metadata = {
            'lines': len(lines),
            'characters': len(content),
            'headings': len([l for l in lines if l.startswith('#')])
        }
        
        return {
            "content": content,
            "metadata": metadata
        }
    
    @classmethod
    def get_supported_formats(cls) -> list:
        """获取支持的文件格式列表"""
        formats = []
        for extensions in cls.SUPPORTED_FORMATS.values():
            formats.extend(extensions)
        return formats


# 便捷函数
def parse_document(file_path: Union[str, Path]) -> Dict:
    """
    便捷的文档解析函数
    
    Args:
        file_path: 文档路径
        
    Returns:
        解析结果字典
    """
    parser = DocumentParser()
    return parser.parse(file_path)


if __name__ == "__main__":
    # 测试代码
    print("文档解析器模块")
    print(f"支持的格式: {DocumentParser.get_supported_formats()}")
    
    parser = DocumentParser()
    print(f"可用的解析器: {parser.available_parsers}")
